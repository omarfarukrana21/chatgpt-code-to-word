import streamlit as st
from docx import Document
import io

st.set_page_config(page_title="ChatGPT Code to Word", layout="centered")

st.title("📄 ChatGPT Code → MS Word Converter")
st.write("ChatGPT থেকে Python code কপি করে Paste করো, তারপর Word ফাইল ডাউনলোড করো।")

code_input = st.text_area(
    "👇 এখানে Python Code Paste করো",
    height=300,
    placeholder="print('Hello from ChatGPT!')"
)

if st.button("Convert to Word & Download"):
    if not code_input.strip():
        st.error("আগে কোড Paste করো 😅")
    else:
        doc = Document()
        doc.add_heading("Python Code from ChatGPT", level=1)
        for line in code_input.split("\n"):
            doc.add_paragraph(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("Word ফাইল তৈরি হয়ে গেছে! নিচের বাটনে ক্লিক করো 👇")
        st.download_button(
            label="📥 Download MS Word File",
            data=buffer,
            file_name="chatgpt_code.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


